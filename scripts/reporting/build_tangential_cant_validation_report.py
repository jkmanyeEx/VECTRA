#!/usr/bin/env python3
"""Build the VECTRA alternating tangential-cant validation report.

The Markdown file is the authoritative prose source. The structured MATLAB
validation JSON is the authoritative numerical source. Generation stops if the
current evidence does not match the validated tangential geometry.
"""

from __future__ import annotations

import importlib.util
import json
import math
import re
import shutil
import tempfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
REPORT_SOURCE = (
    ROOT / "docs" / "reports" / "tangential-cant-yaw-validation-report.md"
)
REFERENCE_DOCX = (
    ROOT / "docs" / "reports" / "cant-model-implementation-validation-report.docx"
)
REPORT_OUTPUT = (
    ROOT / "docs" / "reports" / "tangential-cant-yaw-validation-report.docx"
)
VALIDATION_RESULT = (
    ROOT / "results" / "reports" / "cant-validation" / "validation-report.json"
)
VALIDATION_LOG = (
    ROOT / "results" / "reports" / "cant-validation" / "console.log"
)
ASSET_DIR = ROOT / "docs" / "reports" / "assets"
GEOMETRY_FIGURE = ASSET_DIR / "tangential-cant-geometry.png"
YAW_FIGURE = ASSET_DIR / "tangential-cant-yaw-coefficients.png"
HOVER_FIGURE = ASSET_DIR / "tangential-cant-hover-qualification.png"
BASE_GENERATOR = (
    ROOT / "scripts" / "reporting" / "build_cant_validation_report.py"
)

TIMES_REGULAR = Path(
    "/System/Library/Fonts/Supplemental/Times New Roman.ttf"
)
TIMES_BOLD = Path(
    "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"
)
BATANG_TTC = Path(
    "/Applications/Microsoft Word.app/Contents/Resources/DFonts/batang.ttc"
)

TIMES_NAME = "Times New Roman"
BATANG_NAME = "Batang"

BLUE = "#2E74B5"
DARK_BLUE = "#16324F"
DEEP_BLUE = "#0B2545"
MUTED = "#65717E"
LIGHT_GRAY = "#F2F4F7"
GRID = "#D8DEE6"
GOLD = "#A87400"
WHITE = "#FFFFFF"
INK = "#20262E"


def load_base_helpers():
    """Load the retained report's stable document helper functions."""
    spec = importlib.util.spec_from_file_location(
        "vectra_base_report_generator", BASE_GENERATOR
    )
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load document helpers: {BASE_GENERATOR}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    module.ASCII_FONT = TIMES_NAME
    module.KOREAN_FONT = BATANG_NAME
    module.MONO_FONT = TIMES_NAME
    original_set_run_font = module.set_run_font

    def set_run_font(
        run,
        *,
        ascii_font: str = TIMES_NAME,
        east_asia_font: str = BATANG_NAME,
        size: float | None = None,
        color: str | None = None,
        bold: bool | None = None,
        italic: bool | None = None,
    ) -> None:
        """Apply the requested Latin and Korean fonts to every explicit run.

        The retained helper's default arguments were bound to its historical
        AppleGothic constants when that module was imported. Replacing only the
        module constants therefore does not update those defaults.
        """
        original_set_run_font(
            run,
            ascii_font=ascii_font,
            east_asia_font=east_asia_font,
            size=size,
            color=color,
            bold=bold,
            italic=italic,
        )

    module.set_run_font = set_run_font
    module.REPORT_SOURCE = REPORT_SOURCE
    module.REPORT_OUTPUT = REPORT_OUTPUT
    return module


BASE = load_base_helpers()


def assert_close(actual: float, expected: float, tolerance: float, label: str):
    if not math.isfinite(actual) or abs(actual - expected) > tolerance:
        raise ValueError(
            f"{label} mismatch: actual={actual!r}, expected={expected!r}"
        )


def load_and_validate_evidence() -> dict:
    """Load fresh evidence and enforce every report-critical invariant."""
    if not VALIDATION_RESULT.is_file():
        raise FileNotFoundError(
            f"Run runCantValidationLogged() first: {VALIDATION_RESULT}"
        )
    result = json.loads(VALIDATION_RESULT.read_text(encoding="utf-8"))
    if not VALIDATION_LOG.is_file():
        raise FileNotFoundError(VALIDATION_LOG)
    console = VALIDATION_LOG.read_text(encoding="utf-8", errors="replace")
    completed_runs = list(
        re.finditer(
            r"UNIT_PASSED=(\d+) UNIT_FAILED=(\d+).*?"
            r"CANT_VALIDATION_PASSED=(\d+).*?"
            r"VECTRA_CANT_VALIDATION_COMPLETE=([^\r\n]+)",
            console,
            flags=re.DOTALL,
        )
    )
    if not completed_runs:
        raise ValueError("No completed cant-validation run was found in console.log.")
    latest = completed_runs[-1]
    if latest.group(1) != "19" or latest.group(2) != "0":
        raise ValueError(
            "Latest automated test result is not 19 passed and 0 failed: "
            f"{latest.group(1)} passed, {latest.group(2)} failed"
        )
    if latest.group(3) != "1":
        raise ValueError("Latest logged cant validation did not pass.")
    if latest.group(4).strip() != "28-Jul-2026 09:28:33":
        raise ValueError(
            "Report timestamp does not match the latest completed validation: "
            f"{latest.group(4).strip()}"
        )
    if result.get("passed") is not True:
        raise ValueError("Structured cant validation did not pass.")

    zero = result.get("zeroCant", {})
    ten = result.get("tenDegree", {})
    if zero.get("passed") is not True:
        raise ValueError("Zero-cant regression did not pass.")
    if ten.get("passed") is not True:
        raise ValueError("Tangential-cant integration validation did not pass.")
    if ten.get("geometryId") != "alternating-tangential-cant-10":
        raise ValueError(f"Unexpected geometryId: {ten.get('geometryId')!r}")
    if ten.get("cantType") != "tangential":
        raise ValueError(f"Unexpected cantType: {ten.get('cantType')!r}")
    if ten.get("motorCantAnglesDeg") != [-10, 10, -10, 10]:
        raise ValueError(
            f"Unexpected motor cant angles: {ten.get('motorCantAnglesDeg')!r}"
        )
    if ten.get("yawSignsAligned") is not True:
        raise ValueError("Yaw signs are not aligned.")
    if ten.get("allocationRank") != 4:
        raise ValueError("Tangential allocation is not rank four.")
    if ten.get("steadyHoverAllocationFeasible") is not True:
        raise ValueError("Steady-hover allocation is infeasible.")
    if ten.get("finite") is not True:
        raise ValueError("Tangential simulation contains non-finite values.")
    if ten.get("motorLimitExceeded") is not False:
        raise ValueError("Tangential simulation exceeded a motor limit.")
    if any(ten.get("steadyHoverNegativeDemand", [True])):
        raise ValueError("Steady hover contains negative squared-RPM demand.")
    if any(ten.get("steadyHoverOverSpeedDemand", [True])):
        raise ValueError("Steady hover contains overspeed demand.")

    expected_signs = [-1, 1, -1, 1]
    yaw = [float(value) for value in ten["yawCoefficientsNmPerRpm2"]]
    zero_yaw = [
        float(value) for value in ten["zeroCantYawCoefficientsNmPerRpm2"]
    ]
    if [1 if value > 0 else -1 for value in yaw] != expected_signs:
        raise ValueError(f"Unexpected tangential yaw signs: {yaw!r}")
    if [1 if value > 0 else -1 for value in zero_yaw] != expected_signs:
        raise ValueError(f"Unexpected zero-cant yaw signs: {zero_yaw!r}")

    assert_close(
        float(ten["expectedVerticalForceScale"]),
        math.cos(math.radians(10)),
        1e-12,
        "Expected vertical-force scale",
    )
    assert_close(
        float(ten["actualVerticalForceScale"]),
        math.cos(math.radians(10)),
        1e-12,
        "Actual vertical-force scale",
    )
    assert_close(
        float(ten["yawAuthorityGain"]),
        min(abs(a / b) for a, b in zip(yaw, zero_yaw)),
        1e-12,
        "Yaw-authority gain",
    )
    if float(ten["yawAuthorityGain"]) <= 1:
        raise ValueError("Yaw-authority gain must be greater than one.")
    if float(ten["steadyHoverAllocationResidualNorm"]) > 1e-12:
        raise ValueError("Steady-hover allocation residual is too large.")
    return result


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = TIMES_BOLD if bold else TIMES_REGULAR
    if not path.is_file():
        raise FileNotFoundError(path)
    return ImageFont.truetype(str(path), size=size)


def draw_centered(
    draw: ImageDraw.ImageDraw,
    xy: tuple[float, float],
    text: str,
    text_font: ImageFont.FreeTypeFont,
    fill: str = INK,
):
    bbox = draw.textbbox((0, 0), text, font=text_font)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.text(
        (xy[0] - width / 2, xy[1] - height / 2),
        text,
        font=text_font,
        fill=fill,
    )


def draw_right(
    draw: ImageDraw.ImageDraw,
    xy: tuple[float, float],
    text: str,
    text_font: ImageFont.FreeTypeFont,
    fill: str = INK,
):
    bbox = draw.textbbox((0, 0), text, font=text_font)
    draw.text(
        (xy[0] - (bbox[2] - bbox[0]), xy[1]),
        text,
        font=text_font,
        fill=fill,
    )


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[float, float],
    end: tuple[float, float],
    fill: str,
    width: int = 7,
    head: int = 18,
):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    left = (
        end[0] - head * math.cos(angle - math.pi / 6),
        end[1] - head * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - head * math.cos(angle + math.pi / 6),
        end[1] - head * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=fill)


def draw_blossom(draw: ImageDraw.ImageDraw, center: tuple[int, int]):
    """Small locked top-right research blossom."""
    cx, cy = center
    for index in range(5):
        angle = -math.pi / 2 + index * 2 * math.pi / 5
        px = cx + 14 * math.cos(angle)
        py = cy + 14 * math.sin(angle)
        draw.ellipse(
            [px - 7, py - 7, px + 7, py + 7],
            fill="#DCE9F5",
            outline=BLUE,
            width=2,
        )
    draw.ellipse(
        [cx - 5, cy - 5, cx + 5, cy + 5],
        fill=GOLD,
        outline=DEEP_BLUE,
        width=1,
    )


def chart_header(
    draw: ImageDraw.ImageDraw, title: str, subtitle: str, width: int
):
    draw.text((90, 58), title, font=font(42, True), fill=DARK_BLUE)
    draw.text((90, 116), subtitle, font=font(24), fill=MUTED)
    draw.line([(90, 160), (width - 90, 160)], fill=GRID, width=3)
    draw_blossom(draw, (width - 118, 88))


def build_geometry_figure():
    width, height = 1800, 1080
    image = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(image)
    chart_header(
        draw,
        "Alternating tangential cant geometry",
        "Top view in QUADSIM_BODY_XY_ZUP; solid arrows show horizontal thrust",
        width,
    )

    center = (670, 585)
    arm = 300
    motor_radius = 52
    draw.line(
        [(center[0] - arm - 100, center[1]), (center[0] + arm + 100, center[1])],
        fill="#B8C2CC",
        width=12,
    )
    draw.line(
        [(center[0], center[1] - arm - 100), (center[0], center[1] + arm + 100)],
        fill="#B8C2CC",
        width=12,
    )
    draw.ellipse(
        [
            center[0] - 58,
            center[1] - 58,
            center[0] + 58,
            center[1] + 58,
        ],
        fill="#E9EFF5",
        outline=DEEP_BLUE,
        width=5,
    )
    draw_centered(draw, center, "CG", font(27, True), DEEP_BLUE)

    motors = [
        ("M1", "+X", "-10 deg", (center[0] + arm, center[1]), (0, 1)),
        ("M2", "+Y", "+10 deg", (center[0], center[1] - arm), (-1, 0)),
        ("M3", "-X", "-10 deg", (center[0] - arm, center[1]), (0, -1)),
        ("M4", "-Y", "+10 deg", (center[0], center[1] + arm), (1, 0)),
    ]
    for name, arm_name, angle_label, position, direction in motors:
        x, y = position
        draw.ellipse(
            [x - motor_radius, y - motor_radius, x + motor_radius, y + motor_radius],
            fill="#F7FAFC",
            outline=DEEP_BLUE,
            width=5,
        )
        draw_centered(draw, (x, y - 4), name, font(24, True), DEEP_BLUE)
        draw_centered(draw, (x, y + 34), arm_name, font(18), MUTED)
        arrow_end = (x + direction[0] * 130, y + direction[1] * 130)
        arrow_start = (
            x + direction[0] * motor_radius,
            y + direction[1] * motor_radius,
        )
        draw_arrow(draw, arrow_start, arrow_end, BLUE, width=10, head=26)
        opposite = (x - direction[0] * 88, y - direction[1] * 88)
        for step in range(0, 74, 18):
            a = step / 88
            b = min(1, (step + 9) / 88)
            draw.line(
                [
                    (x + (opposite[0] - x) * a, y + (opposite[1] - y) * a),
                    (x + (opposite[0] - x) * b, y + (opposite[1] - y) * b),
                ],
                fill="#AAB4BF",
                width=4,
            )
        label_position = (
            x + direction[0] * 175,
            y + direction[1] * 175,
        )
        draw_centered(draw, label_position, angle_label, font(22, True), GOLD)

    draw_arrow(
        draw,
        (center[0] + 20, center[1] + 20),
        (center[0] + 160, center[1] + 20),
        DEEP_BLUE,
        width=5,
        head=17,
    )
    draw.text(
        (center[0] + 168, center[1] + 3),
        "+X",
        font=font(22, True),
        fill=DEEP_BLUE,
    )
    draw_arrow(
        draw,
        (center[0] + 20, center[1] + 20),
        (center[0] + 20, center[1] - 120),
        DEEP_BLUE,
        width=5,
        head=17,
    )
    draw.text(
        (center[0] - 2, center[1] - 156),
        "+Y",
        font=font(22, True),
        fill=DEEP_BLUE,
    )

    box = (1120, 255, 1680, 875)
    draw.rounded_rectangle(box, radius=28, fill="#F7F9FB", outline=GRID, width=3)
    draw.text((1180, 310), "Cant-angle definition", font=font(30, True), fill=DARK_BLUE)
    origin = (1395, 705)
    draw_arrow(draw, origin, (1395, 425), DEEP_BLUE, width=8, head=24)
    draw_arrow(draw, origin, (1600, 705), "#8A98A8", width=8, head=24)
    theta = math.radians(10)
    canted_end = (
        origin[0] + 280 * math.sin(theta),
        origin[1] - 280 * math.cos(theta),
    )
    draw_arrow(draw, origin, canted_end, BLUE, width=11, head=28)
    draw.arc(
        [origin[0] - 86, origin[1] - 86, origin[0] + 86, origin[1] + 86],
        260,
        270,
        fill=GOLD,
        width=7,
    )
    draw.text((1418, 538), "10 deg", font=font(24, True), fill=GOLD)
    draw.text((1420, 402), "0 deg: vertical", font=font(22), fill=DEEP_BLUE)
    draw.text((1430, 735), "90 deg: tangential", font=font(22), fill=MUTED)
    draw.text((1455, 483), "u_i", font=font(24, True), fill=BLUE)
    draw.text((1605, 680), "e_t,i", font=font(22), fill=MUTED)

    draw.text(
        (90, height - 74),
        "Source: cant_tangential_10.json and buildRotorGeometry.m",
        font=font(20),
        fill=MUTED,
    )
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image.save(GEOMETRY_FIGURE, dpi=(220, 220))


def build_yaw_figure(result: dict):
    ten = result["tenDegree"]
    zero = [
        float(value) * 1e9
        for value in ten["zeroCantYawCoefficientsNmPerRpm2"]
    ]
    tangent = [
        float(value) * 1e9 for value in ten["yawCoefficientsNmPerRpm2"]
    ]
    width, height = 1800, 1050
    image = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(image)
    chart_header(
        draw,
        "Signed yaw coefficient by motor",
        "Zero cant vs alternating tangential 10 deg; values in x 10^-9 N m / rpm^2",
        width,
    )

    plot = (180, 235, 1660, 830)
    y_min, y_max = -10.0, 10.0

    def y_pos(value: float) -> float:
        return plot[3] - (value - y_min) / (y_max - y_min) * (
            plot[3] - plot[1]
        )

    for tick in [-10, -5, 0, 5, 10]:
        y = y_pos(tick)
        draw.line(
            [(plot[0], y), (plot[2], y)],
            fill=DEEP_BLUE if tick == 0 else "#E3E8ED",
            width=4 if tick == 0 else 2,
        )
        draw_right(draw, (plot[0] - 26, y - 13), f"{tick:+d}", font(22), MUTED)

    motors = ["M1", "M2", "M3", "M4"]
    group_width = (plot[2] - plot[0]) / 4
    bar_width = 82
    for index, motor in enumerate(motors):
        center_x = plot[0] + group_width * (index + 0.5)
        draw_centered(draw, (center_x, plot[3] + 62), motor, font(26, True), DEEP_BLUE)
        for series_index, value in enumerate([zero[index], tangent[index]]):
            x0 = center_x + (-1 if series_index == 0 else 1) * 54 - bar_width / 2
            x1 = x0 + bar_width
            y0 = y_pos(0)
            y1 = y_pos(value)
            top, bottom = min(y0, y1), max(y0, y1)
            if series_index == 0:
                draw.rectangle(
                    [x0, top, x1, bottom],
                    fill="#FFFFFF",
                    outline=GOLD,
                    width=7,
                )
                for yy in range(int(top) + 12, int(bottom), 18):
                    draw.line([(x0 + 8, yy), (x1 - 8, yy)], fill="#D9BC78", width=3)
            else:
                draw.rectangle(
                    [x0, top, x1, bottom],
                    fill=BLUE,
                    outline=DEEP_BLUE,
                    width=3,
                )
            label_y = top - 30 if value > 0 else bottom + 8
            draw_centered(
                draw,
                ((x0 + x1) / 2, label_y),
                f"{value:+.4f}",
                font(20, True),
                DEEP_BLUE,
            )

    legend_y = 198
    draw.rectangle(
        [525, legend_y - 20, 575, legend_y + 20],
        fill=WHITE,
        outline=GOLD,
        width=5,
    )
    draw.text((590, legend_y - 15), "Zero cant", font=font(22), fill=INK)
    draw.rectangle(
        [845, legend_y - 20, 895, legend_y + 20],
        fill=BLUE,
        outline=DEEP_BLUE,
        width=2,
    )
    draw.text(
        (910, legend_y - 15),
        "Tangential 10 deg",
        font=font(22),
        fill=INK,
    )
    gain = float(ten["yawAuthorityGain"])
    draw.rounded_rectangle(
        [1280, 172, 1635, 224],
        radius=20,
        fill="#F4F6F9",
        outline=GRID,
        width=2,
    )
    draw_centered(
        draw,
        (1457, 198),
        f"Magnitude gain: {gain:.4f}x",
        font(24, True),
        DEEP_BLUE,
    )
    image.save(YAW_FIGURE, dpi=(220, 220))


def build_hover_figure(result: dict):
    ten = result["tenDegree"]
    width, height = 1800, 1050
    image = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(image)
    chart_header(
        draw,
        "Hover qualification at tangential 10 deg",
        "Focused normalized scales; zero-cant reference = 1.000000",
        width,
    )

    panels = [
        (
            "Vertical-force scale",
            float(ten["expectedVerticalForceScale"]),
            float(ten["actualVerticalForceScale"]),
            0.975,
            1.005,
            "Lower is less vertical thrust at equal rpm",
        ),
        (
            "Hover RPM ratio",
            float(ten["expectedHoverRpmRatio"]),
            float(ten["actualHoverRpmRatio"]),
            0.995,
            1.015,
            "Higher is more rpm to maintain hover",
        ),
    ]
    panel_boxes = [(105, 235, 855, 880), (945, 235, 1695, 880)]
    for box, panel in zip(panel_boxes, panels):
        title, expected, actual, minimum, maximum, note = panel
        x0, y0, x1, y1 = box
        draw.rounded_rectangle(box, radius=26, fill="#FAFBFC", outline=GRID, width=3)
        draw.text((x0 + 45, y0 + 42), title, font=font(30, True), fill=DARK_BLUE)
        draw.text((x0 + 45, y0 + 91), note, font=font(20), fill=MUTED)

        axis_x0, axis_x1 = x0 + 95, x1 - 70
        axis_y = y0 + 360
        draw.line([(axis_x0, axis_y), (axis_x1, axis_y)], fill="#AAB4BF", width=5)
        for tick_index in range(5):
            tick = minimum + (maximum - minimum) * tick_index / 4
            x = axis_x0 + (axis_x1 - axis_x0) * tick_index / 4
            draw.line([(x, axis_y - 12), (x, axis_y + 12)], fill="#7B8794", width=3)
            draw_centered(draw, (x, axis_y + 45), f"{tick:.3f}", font(19), MUTED)

        def x_pos(value: float) -> float:
            return axis_x0 + (value - minimum) / (maximum - minimum) * (
                axis_x1 - axis_x0
            )

        ref_x = x_pos(1.0)
        draw.line(
            [(ref_x, y0 + 220), (ref_x, axis_y + 80)],
            fill="#9EA8B3",
            width=3,
        )
        draw_centered(draw, (ref_x, y0 + 197), "zero-cant", font(18), MUTED)

        expected_x = x_pos(expected)
        actual_x = x_pos(actual)
        draw.line(
            [(expected_x, axis_y - 92), (actual_x, axis_y + 92)],
            fill=GRID,
            width=5,
        )
        draw.ellipse(
            [expected_x - 17, axis_y - 109, expected_x + 17, axis_y - 75],
            fill=WHITE,
            outline=GOLD,
            width=6,
        )
        draw.polygon(
            [
                (actual_x, axis_y + 68),
                (actual_x + 20, axis_y + 94),
                (actual_x, axis_y + 120),
                (actual_x - 20, axis_y + 94),
            ],
            fill=BLUE,
            outline=DEEP_BLUE,
        )
        draw_centered(
            draw,
            (expected_x, axis_y - 150),
            f"Analytic {expected:.10f}",
            font(21, True),
            GOLD,
        )
        draw_centered(
            draw,
            (actual_x, axis_y + 158),
            f"Simulated {actual:.10f}",
            font(21, True),
            BLUE,
        )
        delta = (actual - expected) * 100
        draw.rounded_rectangle(
            [x0 + 105, y1 - 82, x1 - 105, y1 - 22],
            radius=18,
            fill="#F2F4F7",
            outline=GRID,
            width=2,
        )
        draw_centered(
            draw,
            ((x0 + x1) / 2, y1 - 52),
            f"Simulated - analytic: {delta:+.4f} percentage point",
            font(21, True),
            DEEP_BLUE,
        )

    image.save(HOVER_FIGURE, dpi=(220, 220))


def clear_paragraph(paragraph):
    """Remove paragraph content while preserving paragraph properties."""
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def clear_document_body(document: Document):
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_document(document: Document):
    BASE.configure_styles(document)
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run("VECTRA  |  VALIDATION REPORT")
    BASE.set_run_font(left, size=8.5, color=BASE.MUTED, bold=True)
    paragraph.add_run("\t")
    right = paragraph.add_run("TANGENTIAL CANT / YAW")
    BASE.set_run_font(right, size=8.5, color=BASE.MUTED)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT
    )

    footer = section.footer
    paragraph = footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run("VECTRA 연구팀 · 2026-07-28")
    BASE.set_run_font(left, size=8.5, color=BASE.MUTED)
    paragraph.add_run("\t")
    page_run = paragraph.add_run("Page ")
    BASE.set_run_font(page_run, size=8.5, color=BASE.MUTED)
    BASE.add_field(page_run, "PAGE")
    of_run = paragraph.add_run(" / ")
    BASE.set_run_font(of_run, size=8.5, color=BASE.MUTED)
    BASE.add_field(of_run, "NUMPAGES")
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT
    )


def add_cover(document: Document):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(24)

    kicker = document.add_paragraph(style="Report Kicker")
    run = kicker.add_run("TECHNICAL VALIDATION REPORT  ·  V1.0")
    BASE.set_run_font(run, size=10, color=BASE.GOLD, bold=True)

    title = document.add_paragraph(style="Report Title")
    run = title.add_run(
        "VECTRA 교대 접선방향 Cant 모델\n및 Yaw 권한 검증 보고서"
    )
    BASE.set_run_font(run, size=25, color=BASE.DARK_BLUE, bold=True)

    subtitle = document.add_paragraph(style="Report Subtitle")
    run = subtitle.add_run(
        "QuadSim 기반 교대 10° 로터축, Yaw 모멘트 및 제어 할당 검증"
    )
    BASE.set_run_font(run, size=13, color=BASE.MUTED)

    metadata = [
        ("검증 환경", "MATLAB/Simulink R2026a"),
        ("검증 완료", "2026-07-28 09:28:33 (Asia/Seoul)"),
        ("작성", "VECTRA 연구팀 - 프로그래밍 및 시뮬레이션 파트"),
        ("형상", "alternating-tangential-cant-10"),
        ("범위", "0° upstream 회귀 + 교대 접선방향 10° yaw/hover 검증"),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        label_run = paragraph.add_run(f"{label}: ")
        BASE.set_run_font(
            label_run, size=10.5, color=BASE.DEEP_BLUE, bold=True
        )
        value_run = paragraph.add_run(value)
        BASE.set_run_font(value_run, size=10.5, color="333333")

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(16)
    BASE.set_paragraph_border(
        rule,
        bottom={
            "val": "single",
            "sz": 22,
            "space": 1,
            "color": BASE.BLUE,
        },
    )

    callout = document.add_paragraph(style="Lead Callout")
    run = callout.add_run(
        "최종 판정  |  yaw coefficient 2.9461× · allocation rank 4 · "
        "steady hover feasible · 19/19 tests passed"
    )
    BASE.set_run_font(run, size=11, color=BASE.DEEP_BLUE, bold=True)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(16)
    run = note.add_run(
        "이 보고서는 모델의 정적 yaw effectiveness, hover balance 및 "
        "구현 타당성을 검증한다. Dynamic yaw response와 실제 비행 성능은 "
        "후속 실험 범위다."
    )
    BASE.set_run_font(
        run, size=10, color=BASE.MUTED, italic=True
    )
    document.add_page_break()


def choose_table_widths(headers: list[str]) -> list[int]:
    count = len(headers)
    first = headers[0]
    if count == 2:
        return [2900, 6460]
    if count == 3:
        if first in {"Motor", "검증 그룹", "지표"}:
            return [2750, 4100, 2510]
        return [2600, 3380, 3380]
    if count == 4:
        if headers[1] == "Analytic/expected":
            return [2700, 2600, 2600, 1460]
        if first in {"구분", "검증 그룹"}:
            return [1700, 3840, 1900, 1920]
        return [3000, 2100, 2100, 2160]
    if count == 5:
        return [900, 1800, 1400, 1550, 3710]
    if count == 6:
        return [1100, 1600, 1100, 1400, 2300, 1860]
    base = 9360 // count
    return [base] * (count - 1) + [9360 - base * (count - 1)]


def install_table_width_override():
    def checked_widths(headers: list[str]) -> list[int]:
        widths = choose_table_widths(headers)
        if len(widths) != len(headers) or sum(widths) != 9360:
            raise ValueError(
                f"Invalid table geometry for {headers!r}: {widths!r}"
            )
        return widths

    BASE.choose_table_widths = checked_widths


def postprocess_document(document: Document):
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(("그림 1.", "그림 2.", "그림 3.")):
            paragraph.style = document.styles["Figure Caption"]
            for run in paragraph.runs:
                BASE.set_run_font(
                    run, size=9.5, color=BASE.MUTED, italic=True
                )
        for run in paragraph.runs:
            if run._element.rPr is None:
                BASE.set_run_font(run)

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def build_document():
    markdown = REPORT_SOURCE.read_text(encoding="utf-8")
    if "<!-- body -->" not in markdown:
        raise ValueError("Report Markdown is missing the body marker.")
    body = markdown.split("<!-- body -->", 1)[1].strip()

    with tempfile.TemporaryDirectory(prefix="vectra-tangential-docx-") as temp:
        working_copy = Path(temp) / "working-reference.docx"
        shutil.copy2(REFERENCE_DOCX, working_copy)
        document = Document(working_copy)
        clear_document_body(document)
        configure_document(document)
        install_table_width_override()
        document.core_properties.title = (
            "VECTRA 교대 접선방향 Cant 모델 및 Yaw 권한 검증 보고서"
        )
        document.core_properties.subject = (
            "Alternating tangential cant yaw-effectiveness validation"
        )
        document.core_properties.author = "VECTRA Research Team"
        document.core_properties.keywords = (
            "VECTRA, QuadSim, tangential cant, yaw, allocation, validation"
        )

        bullet_num_id = BASE.add_numbering_definition(document, decimal=False)
        decimal_num_id = BASE.add_numbering_definition(document, decimal=True)
        add_cover(document)
        BASE.parse_markdown_body(
            document, body, bullet_num_id, decimal_num_id
        )
        postprocess_document(document)
        REPORT_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
        document.save(REPORT_OUTPUT)


def main():
    if not REPORT_SOURCE.is_file():
        raise FileNotFoundError(REPORT_SOURCE)
    if not REFERENCE_DOCX.is_file():
        raise FileNotFoundError(REFERENCE_DOCX)
    if not BATANG_TTC.is_file():
        raise FileNotFoundError(
            "Batang is required but was not found in Microsoft Word resources: "
            f"{BATANG_TTC}"
        )
    result = load_and_validate_evidence()
    build_geometry_figure()
    build_yaw_figure(result)
    build_hover_figure(result)
    build_document()
    print(f"WROTE={GEOMETRY_FIGURE}")
    print(f"WROTE={YAW_FIGURE}")
    print(f"WROTE={HOVER_FIGURE}")
    print(f"WROTE={REPORT_OUTPUT}")


if __name__ == "__main__":
    main()
