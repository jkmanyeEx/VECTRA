#!/usr/bin/env python3
"""Build the VECTRA Cant implementation and validation report."""

from __future__ import annotations

import json
import re
from html import escape
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image as RLImage,
    ListFlowable,
    ListItem,
    PageBreak,
    PageTemplate,
    Paragraph,
    Preformatted,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
REPORT_SOURCE = (
    ROOT / "docs" / "reports" / "cant-model-implementation-validation-report.md"
)
REPORT_OUTPUT = (
    ROOT / "docs" / "reports" / "cant-model-implementation-validation-report.docx"
)
PDF_OUTPUT = (
    ROOT / "docs" / "reports" / "cant-model-implementation-validation-report.pdf"
)
ASSET_DIR = ROOT / "docs" / "reports" / "assets"
CHART_OUTPUT = ASSET_DIR / "cant-validation-comparison.png"
VALIDATION_RESULT = (
    ROOT / "results" / "reports" / "cant-validation" / "validation-report.json"
)

ASCII_FONT = "Arial Unicode MS"
KOREAN_FONT = "Arial Unicode MS"
MONO_FONT = "Menlo"
KOREAN_FONT_PATH = Path("/System/Library/Fonts/AppleSDGothicNeo.ttc")
PDF_KOREAN_FONT_PATH = Path(
    "/System/Library/Fonts/Supplemental/AppleGothic.ttf"
)
MONO_FONT_PATH = Path("/System/Library/Fonts/Menlo.ttc")

BLUE = "2E74B5"
DARK_BLUE = "16324F"
DEEP_BLUE = "0B2545"
MUTED = "65717E"
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
GRID = "D8DEE6"
GOLD = "A87400"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(
    run,
    *,
    ascii_font: str = ASCII_FONT,
    east_asia_font: str = KOREAN_FONT,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = ascii_font
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:eastAsia"), east_asia_font)
    fonts.set(qn("w:cs"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float, color: str = "222222", bold=False) -> None:
    style.font.name = ASCII_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    style._element.get_or_add_rPr()
    fonts = style._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), ASCII_FONT)
    fonts.set(qn("w:hAnsi"), ASCII_FONT)
    fonts.set(qn("w:eastAsia"), KOREAN_FONT)
    fonts.set(qn("w:cs"), KOREAN_FONT)


def set_paragraph_border(paragraph, *, bottom=None, left=None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge_name, edge in (("bottom", bottom), ("left", left)):
        if edge is None:
            continue
        element = OxmlElement(f"w:{edge_name}")
        element.set(qn("w:val"), edge.get("val", "single"))
        element.set(qn("w:sz"), str(edge.get("sz", 8)))
        element.set(qn("w:space"), str(edge.get("space", 4)))
        element.set(qn("w:color"), edge.get("color", BLUE))
        borders.append(element)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def set_keep_with_next(paragraph, enabled=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:keepNext"))
    if enabled and existing is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not enabled and existing is not None:
        p_pr.remove(existing)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)


def set_table_borders(table, color=GRID, size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"Table widths must sum to 9360 DXA: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(run, field_name: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_numbering_definition(document: Document, *, decimal: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if decimal else "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if decimal else "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_paragraph(document: Document, text: str, num_id: int):
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([level, number])
    p_pr.insert(0, num_pr)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    add_inline_text(paragraph, text)
    return paragraph


def add_inline_text(paragraph, text: str) -> None:
    token_pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                ascii_font=MONO_FONT,
                east_asia_font=KOREAN_FONT,
                size=9.5,
                color=DEEP_BLUE,
            )
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run)


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph(style="Equation Block")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(
            run,
            ascii_font=MONO_FONT,
            east_asia_font=KOREAN_FONT,
            size=9,
            color=DEEP_BLUE,
        )


def choose_table_widths(headers: list[str]) -> list[int]:
    count = len(headers)
    if count == 2:
        return [2550, 6810]
    if count == 3:
        return [2800, 3280, 3280]
    if count == 4 and headers[0] == "구분":
        return [1450, 3370, 1540, 3000]
    if count == 4:
        return [3200, 2050, 2050, 2060]
    return [9360 // count] * (count - 1) + [
        9360 - (9360 // count) * (count - 1)
    ]


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    headers = rows[0]
    table = document.add_table(rows=len(rows), cols=len(headers))
    widths = choose_table_widths(headers)
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row_index, row_values in enumerate(rows):
        for column_index, value in enumerate(row_values):
            cell = table.cell(row_index, column_index)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            if column_index > 0 and len(value) < 34:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_text(paragraph, value)
            for run in paragraph.runs:
                set_run_font(
                    run,
                    size=9.2,
                    bold=(row_index == 0),
                    color=DEEP_BLUE if row_index == 0 else "222222",
                )
    trailing = document.add_paragraph()
    trailing.paragraph_format.space_after = Pt(2)


def add_callout(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Lead Callout")
    add_inline_text(paragraph, text)


def add_image(document: Document, image_path: Path, alt_text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(6.3))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)


def parse_markdown_body(
    document: Document, markdown_body: str, bullet_num_id: int, decimal_num_id: int
) -> None:
    lines = markdown_body.splitlines()
    index = 0
    paragraph_buffer: list[str] = []
    in_numbered_list = False
    active_decimal_num_id = decimal_num_id

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph = document.add_paragraph()
        add_inline_text(paragraph, " ".join(line.strip() for line in paragraph_buffer))
        paragraph_buffer.clear()

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        numbered_line = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered_line is None:
            in_numbered_list = False

        if stripped == "<!-- pagebreak -->":
            flush_paragraph()
            document.add_page_break()
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(document, code_lines)
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            add_callout(document, " ".join(quote_lines))
            continue

        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            flush_paragraph()
            image_path = REPORT_SOURCE.parent / image_match.group(2)
            add_image(document, image_path, image_match.group(1))
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines):
            separator = lines[index + 1].strip()
            if separator.startswith("|") and re.fullmatch(
                r"\|[\s:|-]+\|", separator
            ):
                flush_paragraph()
                table_rows = []
                table_rows.append(
                    [cell.strip() for cell in stripped.strip("|").split("|")]
                )
                index += 2
                while index < len(lines) and lines[index].strip().startswith("|"):
                    table_rows.append(
                        [
                            cell.strip()
                            for cell in lines[index].strip().strip("|").split("|")
                        ]
                    )
                    index += 1
                add_markdown_table(document, table_rows)
                continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            if level == 1:
                index += 1
                continue
            paragraph = document.add_paragraph(
                style="Heading 1" if level == 2 else "Heading 2"
            )
            add_inline_text(paragraph, heading.group(2))
            set_keep_with_next(paragraph)
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            flush_paragraph()
            item_parts = [bullet.group(1)]
            index += 1
            while index < len(lines):
                continuation = lines[index].strip()
                if (
                    not continuation
                    or re.match(r"^(?:-|\d+\.)\s+", continuation)
                    or continuation.startswith("#")
                ):
                    break
                item_parts.append(continuation)
                index += 1
            add_list_paragraph(
                document, " ".join(item_parts), bullet_num_id
            )
            continue

        if numbered_line:
            flush_paragraph()
            if not in_numbered_list:
                active_decimal_num_id = add_numbering_definition(
                    document, decimal=True
                )
            item_parts = [numbered_line.group(1)]
            in_numbered_list = True
            index += 1
            while index < len(lines):
                continuation = lines[index].strip()
                if (
                    not continuation
                    or re.match(r"^(?:-|\d+\.)\s+", continuation)
                    or continuation.startswith("#")
                ):
                    break
                item_parts.append(continuation)
                index += 1
            add_list_paragraph(
                document, " ".join(item_parts), active_decimal_num_id
            )
            continue

        if stripped == "":
            flush_paragraph()
            index += 1
            continue

        paragraph_buffer.append(raw)
        index += 1

    flush_paragraph()


def pil_font(size: int, *, bold=False) -> ImageFont.FreeTypeFont:
    path = KOREAN_FONT_PATH
    if not path.exists():
        return ImageFont.load_default()
    index = 7 if bold else 3
    try:
        return ImageFont.truetype(str(path), size=size, index=index)
    except OSError:
        return ImageFont.truetype(str(path), size=size)


def draw_text_centered(draw, box, text, font, fill) -> None:
    left, top, right, bottom = box
    bbox = draw.textbbox((0, 0), text, font=font)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.text(
        ((left + right - width) / 2, (top + bottom - height) / 2),
        text,
        font=font,
        fill=fill,
    )


def draw_dot_panel(
    draw: ImageDraw.ImageDraw,
    *,
    box: tuple[int, int, int, int],
    title: str,
    subtitle: str,
    values: list[tuple[str, float, float]],
    domain: tuple[float, float],
    ticks: list[float],
) -> None:
    left, top, right, bottom = box
    draw.rounded_rectangle(box, radius=28, fill="#FBFCFE", outline="#D8DEE6", width=3)
    draw.text((left + 50, top + 38), title, font=pil_font(38, bold=True), fill="#16324F")
    draw.text((left + 50, top + 92), subtitle, font=pil_font(24), fill="#65717E")

    plot_left, plot_right = left + 160, right - 70
    plot_top, plot_bottom = top + 200, bottom - 115
    d_min, d_max = domain

    def x_of(value: float) -> float:
        return plot_left + (value - d_min) / (d_max - d_min) * (
            plot_right - plot_left
        )

    for tick in ticks:
        x = x_of(tick)
        draw.line((x, plot_top, x, plot_bottom), fill="#E5E9EF", width=2)
        label = f"{tick:.3f}"
        bbox = draw.textbbox((0, 0), label, font=pil_font(20))
        draw.text(
            (x - (bbox[2] - bbox[0]) / 2, plot_bottom + 18),
            label,
            font=pil_font(20),
            fill="#65717E",
        )

    reference_x = x_of(1.0)
    draw.line(
        (reference_x, plot_top - 10, reference_x, plot_bottom),
        fill="#16324F",
        width=3,
    )
    draw.text(
        (reference_x - 45, plot_top - 42),
        "기준 1.0",
        font=pil_font(19, bold=True),
        fill="#16324F",
    )

    row_positions = [plot_top + 95, plot_top + 290]
    for row_index, (label, expected, actual) in enumerate(values):
        y = row_positions[row_index]
        draw.text(
            (left + 48, y - 18),
            label,
            font=pil_font(26, bold=True),
            fill="#222222",
        )
        draw.line((plot_left, y, plot_right, y), fill="#D8DEE6", width=2)

        expected_x = x_of(expected)
        actual_x = x_of(actual)
        expected_y = y - 26
        actual_y = y + 26
        radius = 12
        draw.ellipse(
            (
                expected_x - radius,
                expected_y - radius,
                expected_x + radius,
                expected_y + radius,
            ),
            fill="#A87400",
            outline="#7A5400",
            width=3,
        )
        draw.rectangle(
            (
                actual_x - radius,
                actual_y - radius,
                actual_x + radius,
                actual_y + radius,
            ),
            fill="#2E74B5",
            outline="#16324F",
            width=3,
        )
        draw.text(
            (expected_x + 18, expected_y - 17),
            f"{expected:.9f}",
            font=pil_font(19),
            fill="#7A5400",
        )
        draw.text(
            (actual_x + 18, actual_y - 17),
            f"{actual:.9f}",
            font=pil_font(19),
            fill="#16324F",
        )


def build_chart(result: dict) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (2200, 1180), "white")
    draw = ImageDraw.Draw(image)
    draw.text(
        (90, 58),
        "0° 및 10° Cant 검증 결과",
        font=pil_font(54, bold=True),
        fill="#16324F",
    )
    draw.text(
        (90, 127),
        "이론 기준과 MATLAB/Simulink 결과 비교 · 1초 smoke-hover 검증",
        font=pil_font(28),
        fill="#65717E",
    )

    legend_y = 184
    draw.ellipse((90, legend_y, 112, legend_y + 22), fill="#A87400")
    draw.text((126, legend_y - 7), "이론/기준", font=pil_font(22), fill="#4B5563")
    draw.rectangle((280, legend_y, 302, legend_y + 22), fill="#2E74B5")
    draw.text((316, legend_y - 7), "시뮬레이션", font=pil_font(22), fill="#4B5563")

    ten = result["tenDegree"]
    draw_dot_panel(
        draw,
        box=(70, 245, 1070, 1040),
        title="수직 추력 비율",
        subtitle="0° 기준 대비 Fz 계수 · 1.0 주변 확대 축",
        values=[
            ("0°", 1.0, 1.0),
            (
                "10°",
                ten["expectedVerticalForceScale"],
                ten["actualVerticalForceScale"],
            ),
        ],
        domain=(0.980, 1.002),
        ticks=[0.980, 0.985, 0.990, 0.995, 1.000],
    )
    draw_dot_panel(
        draw,
        box=(1130, 245, 2130, 1040),
        title="호버 RPM 비율",
        subtitle="0° 기준 대비 steady-window 평균 · 1.0 주변 확대 축",
        values=[
            ("0°", 1.0, 1.0),
            (
                "10°",
                ten["expectedHoverRpmRatio"],
                ten["actualHoverRpmRatio"],
            ),
        ],
        domain=(0.998, 1.010),
        ticks=[0.998, 1.000, 1.002, 1.004, 1.006, 1.008, 1.010],
    )
    draw.text(
        (90, 1100),
        "Source: VECTRA validation-report.json · 2026-07-27 · 확대 축은 작은 변화 비교용",
        font=pil_font(21),
        fill="#65717E",
    )
    image.save(CHART_OUTPUT, "PNG", optimize=True)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    heading1 = styles["Heading 1"]
    set_style_font(heading1, 16, BLUE, True)
    heading1.paragraph_format.space_before = Pt(16)
    heading1.paragraph_format.space_after = Pt(8)
    heading1.paragraph_format.keep_with_next = True

    heading2 = styles["Heading 2"]
    set_style_font(heading2, 13, BLUE, True)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.keep_with_next = True

    custom_styles = {
        "Report Kicker": (10, GOLD, True, 0, 8),
        "Report Title": (25, DARK_BLUE, True, 0, 6),
        "Report Subtitle": (14, MUTED, False, 0, 18),
        "Lead Callout": (11, DEEP_BLUE, True, 8, 8),
        "Equation Block": (9, DEEP_BLUE, False, 5, 8),
        "Figure Caption": (9.5, MUTED, False, 2, 8),
    }
    for name, (size, color, bold, before, after) in custom_styles.items():
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style, size, color, bold)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.08

    lead = styles["Lead Callout"]
    set_paragraph_shading_proxy(lead, CALLOUT_FILL)
    set_paragraph_border_proxy(
        lead, left={"val": "single", "sz": 18, "space": 8, "color": BLUE}
    )
    lead.paragraph_format.left_indent = Inches(0.12)
    lead.paragraph_format.right_indent = Inches(0.12)

    equation = styles["Equation Block"]
    set_paragraph_shading_proxy(equation, "F6F8FA")
    set_paragraph_border_proxy(
        equation, left={"val": "single", "sz": 12, "space": 8, "color": GRID}
    )
    equation.paragraph_format.left_indent = Inches(0.12)
    equation.paragraph_format.right_indent = Inches(0.12)


def set_paragraph_shading_proxy(style, fill: str) -> None:
    p_pr = style._element.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def set_paragraph_border_proxy(style, *, left=None) -> None:
    p_pr = style._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    if left is not None:
        node = OxmlElement("w:left")
        node.set(qn("w:val"), left.get("val", "single"))
        node.set(qn("w:sz"), str(left.get("sz", 8)))
        node.set(qn("w:space"), str(left.get("space", 4)))
        node.set(qn("w:color"), left.get("color", BLUE))
        borders.append(node)
    p_pr.append(borders)


def configure_section(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run("VECTRA  |  TECHNICAL VALIDATION REPORT")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    paragraph.add_run("\t")
    right = paragraph.add_run("CANT-ANGLE MVP")
    set_run_font(right, size=8.5, color=MUTED)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT
    )

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run("VECTRA 연구팀 · 2026-07-27")
    set_run_font(left, size=8.5, color=MUTED)
    paragraph.add_run("\t")
    page_run = paragraph.add_run("Page ")
    set_run_font(page_run, size=8.5, color=MUTED)
    add_field(page_run, "PAGE")
    of_run = paragraph.add_run(" / ")
    set_run_font(of_run, size=8.5, color=MUTED)
    add_field(of_run, "NUMPAGES")
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT
    )


def add_cover(document: Document) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(26)

    kicker = document.add_paragraph(style="Report Kicker")
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = kicker.add_run("TECHNICAL VALIDATION REPORT  ·  V1.0")
    set_run_font(run, size=10, color=GOLD, bold=True)

    title = document.add_paragraph(style="Report Title")
    run = title.add_run("VECTRA Cant-Angle 모델\n구현 및 검증 보고서")
    set_run_font(run, size=25, color=DARK_BLUE, bold=True)

    subtitle = document.add_paragraph(style="Report Subtitle")
    run = subtitle.add_run(
        "QuadSim 기반 3차원 로터 추력축, 동역학 및 제어 할당 구현"
    )
    set_run_font(run, size=14, color=MUTED)

    metadata = [
        ("검증 환경", "MATLAB/Simulink R2026a"),
        ("검증 완료", "2026-07-27 13:00:43 (Asia/Seoul)"),
        ("작성", "VECTRA 연구팀 - 프로그래밍 및 시뮬레이션 파트"),
        ("범위", "0° upstream 회귀 + 10° radial-outward 방향성 검증"),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=DEEP_BLUE, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color="333333")

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(18)
    set_paragraph_border(
        rule, bottom={"val": "single", "sz": 22, "space": 1, "color": BLUE}
    )

    callout = document.add_paragraph(style="Lead Callout")
    run = callout.add_run(
        "최종 판정  |  단위 테스트 9개 통과 · 0° 회귀 통과 · "
        "10° 물리 방향성 검증 통과 · overall passed = true"
    )
    set_run_font(run, size=11, color=DEEP_BLUE, bold=True)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "이 보고서는 Cant 구현의 타당성을 검증한다. 실제 기체의 최적 cant "
        "angle이나 효율 향상은 후속 다요인 실험의 연구 질문이다."
    )
    set_run_font(run, size=10, color=MUTED, italic=True)
    document.add_page_break()


def postprocess_caption_styles(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("그림 1."):
            paragraph.style = document.styles["Figure Caption"]
            for run in paragraph.runs:
                set_run_font(run, size=9.5, color=MUTED, italic=True)
            set_keep_with_next(paragraph, False)


def build_document(markdown_body: str) -> None:
    document = Document()
    configure_styles(document)
    configure_section(document)
    document.core_properties.title = "VECTRA Cant-Angle 모델 구현 및 검증 보고서"
    document.core_properties.subject = "QuadSim Cant-Angle MVP technical validation"
    document.core_properties.author = "VECTRA Research Team"
    document.core_properties.keywords = (
        "VECTRA, QuadSim, Cant Angle, MATLAB, Simulink, Quadcopter"
    )

    bullet_num_id = add_numbering_definition(document, decimal=False)
    decimal_num_id = add_numbering_definition(document, decimal=True)

    add_cover(document)
    parse_markdown_body(document, markdown_body, bullet_num_id, decimal_num_id)
    postprocess_caption_styles(document)

    REPORT_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(REPORT_OUTPUT)


def pdf_markup(text: str) -> str:
    escaped = escape(text)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(
        r"`(.+?)`",
        r'<font name="Courier" color="#0B2545" size="8.8">\1</font>',
        escaped,
    )
    return escaped


def build_pdf_styles() -> dict[str, ParagraphStyle]:
    if "AppleGothic" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("AppleGothic", str(PDF_KOREAN_FONT_PATH)))
        pdfmetrics.registerFontFamily(
            "AppleGothic",
            normal="AppleGothic",
            bold="AppleGothic",
            italic="AppleGothic",
            boldItalic="AppleGothic",
        )
    sample = getSampleStyleSheet()
    return {
        "body": ParagraphStyle(
            "VECTRA Body",
            parent=sample["BodyText"],
            fontName="AppleGothic",
            fontSize=10.2,
            leading=15.2,
            textColor=colors.HexColor("#222222"),
            spaceAfter=6,
            wordWrap="CJK",
            allowWidows=0,
            allowOrphans=0,
        ),
        "h1": ParagraphStyle(
            "VECTRA H1",
            parent=sample["Heading1"],
            fontName="AppleGothic",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor(f"#{BLUE}"),
            spaceBefore=16,
            spaceAfter=8,
            keepWithNext=True,
            wordWrap="CJK",
        ),
        "h2": ParagraphStyle(
            "VECTRA H2",
            parent=sample["Heading2"],
            fontName="AppleGothic",
            fontSize=13,
            leading=17,
            textColor=colors.HexColor(f"#{BLUE}"),
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=True,
            wordWrap="CJK",
        ),
        "callout": ParagraphStyle(
            "VECTRA Callout",
            parent=sample["BodyText"],
            fontName="AppleGothic",
            fontSize=10.2,
            leading=15,
            textColor=colors.HexColor(f"#{DEEP_BLUE}"),
            backColor=colors.HexColor(f"#{CALLOUT_FILL}"),
            borderColor=colors.HexColor(f"#{BLUE}"),
            borderWidth=1.2,
            borderPadding=9,
            leftIndent=5,
            rightIndent=5,
            spaceBefore=6,
            spaceAfter=10,
            wordWrap="CJK",
        ),
        "caption": ParagraphStyle(
            "VECTRA Caption",
            parent=sample["BodyText"],
            fontName="AppleGothic",
            fontSize=8.8,
            leading=12,
            textColor=colors.HexColor(f"#{MUTED}"),
            alignment=TA_LEFT,
            spaceAfter=9,
            wordWrap="CJK",
        ),
        "table": ParagraphStyle(
            "VECTRA Table",
            parent=sample["BodyText"],
            fontName="AppleGothic",
            fontSize=8.2,
            leading=11,
            textColor=colors.HexColor("#222222"),
            wordWrap="CJK",
        ),
        "table_header": ParagraphStyle(
            "VECTRA Table Header",
            parent=sample["BodyText"],
            fontName="AppleGothic",
            fontSize=8.2,
            leading=11,
            textColor=colors.HexColor(f"#{DEEP_BLUE}"),
            wordWrap="CJK",
        ),
        "code": ParagraphStyle(
            "VECTRA Code",
            fontName="Courier",
            fontSize=8,
            leading=10.5,
            textColor=colors.HexColor(f"#{DEEP_BLUE}"),
            backColor=colors.HexColor("#F6F8FA"),
            borderColor=colors.HexColor(f"#{GRID}"),
            borderWidth=0.6,
            borderPadding=7,
            leftIndent=6,
            rightIndent=6,
            spaceBefore=4,
            spaceAfter=8,
        ),
        "cover_kicker": ParagraphStyle(
            "VECTRA Cover Kicker",
            fontName="AppleGothic",
            fontSize=9.5,
            leading=12,
            textColor=colors.HexColor(f"#{GOLD}"),
            spaceAfter=9,
        ),
        "cover_title": ParagraphStyle(
            "VECTRA Cover Title",
            fontName="AppleGothic",
            fontSize=25,
            leading=33,
            textColor=colors.HexColor(f"#{DARK_BLUE}"),
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "cover_subtitle": ParagraphStyle(
            "VECTRA Cover Subtitle",
            fontName="AppleGothic",
            fontSize=13,
            leading=18,
            textColor=colors.HexColor(f"#{MUTED}"),
            spaceAfter=19,
            wordWrap="CJK",
        ),
        "cover_meta": ParagraphStyle(
            "VECTRA Cover Meta",
            fontName="AppleGothic",
            fontSize=9.8,
            leading=14,
            textColor=colors.HexColor("#333333"),
            spaceAfter=3,
            wordWrap="CJK",
        ),
    }


def pdf_page_furniture(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("AppleGothic", 7.5)
    canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
    canvas.drawString(
        inch,
        letter[1] - 0.46 * inch,
        "VECTRA  |  TECHNICAL VALIDATION REPORT",
    )
    canvas.drawRightString(
        letter[0] - inch,
        letter[1] - 0.46 * inch,
        "CANT-ANGLE MVP",
    )
    canvas.drawString(inch, 0.48 * inch, "VECTRA 연구팀 · 2026-07-27")
    canvas.drawRightString(
        letter[0] - inch, 0.48 * inch, f"Page {canvas.getPageNumber()}"
    )
    canvas.restoreState()


def add_pdf_cover(story: list, styles: dict[str, ParagraphStyle]) -> None:
    story.append(Spacer(1, 0.45 * inch))
    story.append(
        Paragraph(
            "TECHNICAL VALIDATION REPORT  ·  V1.0",
            styles["cover_kicker"],
        )
    )
    story.append(
        Paragraph(
            "VECTRA Cant-Angle 모델<br/>구현 및 검증 보고서",
            styles["cover_title"],
        )
    )
    story.append(
        Paragraph(
            "QuadSim 기반 3차원 로터 추력축, 동역학 및 제어 할당 구현",
            styles["cover_subtitle"],
        )
    )
    metadata = [
        ("검증 환경", "MATLAB/Simulink R2026a"),
        ("검증 완료", "2026-07-27 13:00:43 (Asia/Seoul)"),
        ("작성", "VECTRA 연구팀 - 프로그래밍 및 시뮬레이션 파트"),
        ("범위", "0° upstream 회귀 + 10° radial-outward 방향성 검증"),
    ]
    for label, value in metadata:
        story.append(
            Paragraph(
                f"<b>{escape(label)}:</b> {escape(value)}",
                styles["cover_meta"],
            )
        )
    story.append(Spacer(1, 0.12 * inch))
    rule = Table([[""]], colWidths=[6.5 * inch], rowHeights=[0.045 * inch])
    rule.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{BLUE}")),
                ("BOX", (0, 0), (-1, -1), 0, colors.HexColor(f"#{BLUE}")),
            ]
        )
    )
    story.append(rule)
    story.append(Spacer(1, 0.18 * inch))
    story.append(
        Paragraph(
            "<b>최종 판정</b>  |  단위 테스트 9개 통과 · 0° 회귀 통과 · "
            "10° 물리 방향성 검증 통과 · overall passed = true",
            styles["callout"],
        )
    )
    story.append(Spacer(1, 0.12 * inch))
    story.append(
        Paragraph(
            "이 보고서는 Cant 구현의 타당성을 검증한다. 실제 기체의 최적 "
            "cant angle이나 효율 향상은 후속 다요인 실험의 연구 질문이다.",
            styles["caption"],
        )
    )
    story.append(PageBreak())


def pdf_table_flowable(
    rows: list[list[str]], styles: dict[str, ParagraphStyle]
) -> Table:
    headers = rows[0]
    widths_dxa = choose_table_widths(headers)
    col_widths = [width / 1440 * inch for width in widths_dxa]
    data = []
    for row_index, row in enumerate(rows):
        row_style = styles["table_header"] if row_index == 0 else styles["table"]
        data.append([Paragraph(pdf_markup(value), row_style) for value in row])
    table = Table(
        data,
        colWidths=col_widths,
        repeatRows=1,
        hAlign="LEFT",
        splitByRow=1,
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{LIGHT_GRAY}")),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor(f"#{GRID}")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def parse_pdf_markdown(
    markdown_body: str, styles: dict[str, ParagraphStyle]
) -> list:
    lines = markdown_body.splitlines()
    story: list = []
    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        text = " ".join(line.strip() for line in paragraph_buffer)
        style = styles["caption"] if text.startswith("**그림 ") else styles["body"]
        story.append(Paragraph(pdf_markup(text), style))
        paragraph_buffer.clear()

    while index < len(lines):
        stripped = lines[index].strip()

        if stripped == "<!-- pagebreak -->":
            flush_paragraph()
            story.append(PageBreak())
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            story.append(Preformatted("\n".join(code_lines), styles["code"]))
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            story.append(
                Paragraph(pdf_markup(" ".join(quote_lines)), styles["callout"])
            )
            continue

        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            flush_paragraph()
            image_path = REPORT_SOURCE.parent / image_match.group(2)
            image = RLImage(str(image_path), width=6.3 * inch, height=3.38 * inch)
            image.hAlign = "CENTER"
            story.extend([Spacer(1, 5), image, Spacer(1, 3)])
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines):
            separator = lines[index + 1].strip()
            if separator.startswith("|") and re.fullmatch(
                r"\|[\s:|-]+\|", separator
            ):
                flush_paragraph()
                rows = [[cell.strip() for cell in stripped.strip("|").split("|")]]
                index += 2
                while index < len(lines) and lines[index].strip().startswith("|"):
                    rows.append(
                        [
                            cell.strip()
                            for cell in lines[index].strip().strip("|").split("|")
                        ]
                    )
                    index += 1
                story.extend(
                    [pdf_table_flowable(rows, styles), Spacer(1, 0.08 * inch)]
                )
                continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            if level > 1:
                story.append(
                    Paragraph(
                        pdf_markup(heading.group(2)),
                        styles["h1"] if level == 2 else styles["h2"],
                    )
                )
            index += 1
            continue

        if re.match(r"^-\s+(.+)$", stripped):
            flush_paragraph()
            items = []
            while index < len(lines):
                match = re.match(r"^-\s+(.+)$", lines[index].strip())
                if not match:
                    break
                parts = [match.group(1)]
                index += 1
                while index < len(lines):
                    continuation = lines[index].strip()
                    if (
                        not continuation
                        or re.match(r"^(?:-|\d+\.)\s+", continuation)
                        or continuation.startswith("#")
                    ):
                        break
                    parts.append(continuation)
                    index += 1
                items.append(
                    ListItem(
                        Paragraph(pdf_markup(" ".join(parts)), styles["body"]),
                        leftIndent=14,
                    )
                )
            story.append(
                ListFlowable(
                    items,
                    bulletType="bullet",
                    start="circle",
                    leftIndent=18,
                    bulletFontName="AppleGothic",
                    bulletFontSize=8,
                    spaceAfter=4,
                )
            )
            continue

        if re.match(r"^\d+\.\s+(.+)$", stripped):
            flush_paragraph()
            items = []
            while index < len(lines):
                match = re.match(r"^\d+\.\s+(.+)$", lines[index].strip())
                if not match:
                    break
                parts = [match.group(1)]
                index += 1
                while index < len(lines):
                    continuation = lines[index].strip()
                    if (
                        not continuation
                        or re.match(r"^(?:-|\d+\.)\s+", continuation)
                        or continuation.startswith("#")
                    ):
                        break
                    parts.append(continuation)
                    index += 1
                items.append(
                    ListItem(
                        Paragraph(pdf_markup(" ".join(parts)), styles["body"]),
                        leftIndent=16,
                    )
                )
            story.append(
                ListFlowable(
                    items,
                    bulletType="1",
                    start="1",
                    leftIndent=22,
                    bulletFontName="AppleGothic",
                    bulletFontSize=9,
                    spaceAfter=4,
                )
            )
            continue

        if stripped == "":
            flush_paragraph()
            index += 1
            continue

        paragraph_buffer.append(lines[index])
        index += 1

    flush_paragraph()
    return story


def build_pdf(markdown_body: str) -> None:
    styles = build_pdf_styles()
    doc = BaseDocTemplate(
        str(PDF_OUTPUT),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.78 * inch,
        bottomMargin=0.72 * inch,
        title="VECTRA Cant-Angle 모델 구현 및 검증 보고서",
        author="VECTRA Research Team",
        subject="QuadSim Cant-Angle MVP technical validation",
    )
    frame = Frame(
        doc.leftMargin,
        doc.bottomMargin,
        doc.width,
        doc.height,
        id="normal",
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )
    doc.addPageTemplates(
        [PageTemplate(id="VECTRA", frames=[frame], onPage=pdf_page_furniture)]
    )
    story: list = []
    add_pdf_cover(story, styles)
    story.extend(parse_pdf_markdown(markdown_body, styles))
    doc.build(story)


def main() -> None:
    if not REPORT_SOURCE.is_file():
        raise FileNotFoundError(REPORT_SOURCE)
    if not VALIDATION_RESULT.is_file():
        raise FileNotFoundError(
            "Run runCantValidationLogged() before building the report: "
            f"{VALIDATION_RESULT}"
        )
    validation = json.loads(VALIDATION_RESULT.read_text(encoding="utf-8"))
    if not validation.get("passed", False):
        raise RuntimeError("Refusing to publish a report for a failed validation.")

    build_chart(validation)
    markdown = REPORT_SOURCE.read_text(encoding="utf-8")
    marker = "<!-- body -->"
    if marker not in markdown:
        raise RuntimeError(f"Missing report body marker: {marker}")
    body = markdown.split(marker, 1)[1]
    build_document(body)
    build_pdf(body)
    print(f"REPORT_DOCX={REPORT_OUTPUT}")
    print(f"REPORT_PDF={PDF_OUTPUT}")
    print(f"REPORT_CHART={CHART_OUTPUT}")


if __name__ == "__main__":
    main()
